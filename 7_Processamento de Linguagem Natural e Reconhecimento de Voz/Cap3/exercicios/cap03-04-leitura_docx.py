# Extraindo Texto de Arquivos DOCX

# Imports
# pip install python-docx
import docx

# Função para obter texto de arquivos DOCX
def getTextWord(wordFileName):
    doc = docx.Document(wordFileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# Nome do arquivo
docName = 'arquivo1.docx'
print('Documento completo:\n', getTextWord(docName))
print("\n")
doc = docx.Document(docName)
print('Número de Parágrafos :', len(doc.paragraphs)-1)
print('Parágrafo 1:', doc.paragraphs[1].text)
print('Número de execuções no parágrafo 1:', len(doc.paragraphs[1].runs))

print("\n")

print('Parágrafo 2:', doc.paragraphs[2].text)
print('Parágrafo 2 estilo:', doc.paragraphs[2].style)

print("\n")

for idx, run in enumerate(doc.paragraphs[1].runs):
    print('Run %s : %s' %(idx,run.text))

print("\n")
print('is Run 3 bold:',doc.paragraphs[1].runs[3].bold)
print('is Run 5 italic:',doc.paragraphs[1].runs[5].italic)
print('is Run 7 underlined:',doc.paragraphs[1].runs[7].underline)